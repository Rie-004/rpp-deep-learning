from docx import Document
from datetime import date

def generate_docx_rpp(data):
    doc = Document()
    doc.add_heading('RPP Deep Learning - Kurikulum Merdeka SD', 0)

    doc.add_heading('Identitas', level=1)
    for k, v in data['identitas'].items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading('1. Identifikasi', level=1)
    for k, v in data['identifikasi'].items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading('2. Desain Pembelajaran', level=1)
    for k, v in data['desain'].items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading('3. Pengalaman Belajar', level=1)
    for k, v in data['pengalaman'].items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading('4. Asesmen Pembelajaran', level=1)
    for k, v in data['asesmen'].items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading('5. Sumber Belajar', level=1)
    doc.add_paragraph(data['sumber'])

    doc.add_heading('6. Refleksi', level=1)
    doc.add_paragraph(data['refleksi'])

    doc.add_paragraph(f"© {date.today().year} Candra Sudirno")
    return doc
